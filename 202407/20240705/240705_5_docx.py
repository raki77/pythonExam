#워드 라이브러리 설치 :  pip install python-docx
# <워드 파일 읽기>

from docx import Document


# 워드 파일 열기
doc = Document('송장.docx')
print("doc type:", type(doc))

# 단락 읽기  
for index, paragraph in enumerate(doc.paragraphs):
  print(f"단락 번호 {index}:") # 단락 인덱스 출력
  # 각 단락 내에서 개별 텍스트 런의 서식을 확인합니다.
  for index, run in enumerate(paragraph.runs):
    print(f"  텍스트 번호: {index}")
    print(f"  텍스트 내용: {run.text}")
    print("  텍스트 번호: " + str(index))

# 표 확인하기
# 문서에 있는 모든 표를 순회
for table_index, table in enumerate(doc.tables):
  print(f"표 {table_index}:") #표의 인덱스 출력
  # 표의 각 행을 순회
  for row_index, row in enumerate(table.rows):
    print(f" 행 {row_index}:") # 행 번호 출력
    # 행의 각 셀을 순회
    for cell_index, cell in enumerate(row.cells):
      print(f"  셀 {cell_index}: {cell.text}") # 셀의 인덱스 출력
  print('-' * 20) # 단락 간 구분선

