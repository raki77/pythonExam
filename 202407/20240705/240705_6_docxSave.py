# < 송장 만들기 >

from docx import Document


# 워드 문서 불러오기
doc = Document('송장.docx')
# 고객 정보를 입력해야하는 단락 지정
target_paragraph = doc.paragraphs[8]
# 각 라벨 뒤에 텍스트 추가
target_paragraph.runs[1].add_text('이호연')
# target_paragraph.runs[2].add_text('서울시 마포구 서교동')
# target_paragraph.runs[4].add_text('070-333-4444')
# 구매 내역을 입력할 표 지정
table = doc.tables[0]
# 표에 입력할 데이터를 'data' 변수에 리스트로 저장
data = [["제품1", "2", "12,000","24,000"]]
# 헤더를 제외한 표의 두 번째 행부터 각 셀에 data를 입력
for i, row_data in enumerate(data):
  for j, cell_data in enumerate(row_data):
    table.cell(i+1, j).text = cell_data
# 표의 마지막 행의 마지막 셀에 '총 합계' 값을 입력
table.rows[-1].cells[-1].text = '240,000'
# 문서 저장
doc.save('이호연_송장.docx')
